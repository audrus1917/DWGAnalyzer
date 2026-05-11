from pathlib import Path

from docx import Document
from openpyxl import Workbook

from parsedwg.docs_ingest import _extract_csv_text, _extract_docx_text, _extract_xlsx_text


def test_extract_csv_text_includes_rows(tmp_path: Path) -> None:
    path = tmp_path / "spec.csv"
    path.write_text("Наименование;Кол-во\nСветильник;12\n", encoding="utf-8")

    text = _extract_csv_text(path)

    assert "Наименование | Кол-во" in text
    assert "Светильник | 12" in text


def test_extract_docx_text_includes_paragraphs_and_table(tmp_path: Path) -> None:
    path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("Первая строка")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Кабель"
    table.rows[0].cells[1].text = "120 м"
    doc.save(path)

    text = _extract_docx_text(path)

    assert "Первая строка" in text
    assert "Кабель | 120 м" in text


def test_extract_xlsx_text_includes_sheet_name_and_cells(tmp_path: Path) -> None:
    path = tmp_path / "spec.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "СО"
    sheet.append(["Наименование", "Кол-во"])
    sheet.append(["Светильник", 12])
    workbook.save(path)

    text = _extract_xlsx_text(path)

    assert "# СО" in text
    assert "Наименование | Кол-во" in text
    assert "Светильник | 12" in text
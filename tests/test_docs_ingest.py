from pathlib import Path

from docx import Document
from openpyxl import Workbook

from parsedwg.docs_ingest import (
    _discover_documents,
    _extract_csv_text,
    _extract_docx_text,
    _extract_glossary_terms_from_pages,
    _extract_xlsx_text,
    _compute_md5_hex,
)


def test_discover_documents_recursively_collects_supported_files(tmp_path: Path) -> None:
    root = tmp_path / "docs"
    nested = root / "nested"
    nested.mkdir(parents=True)

    (root / "a.docx").write_bytes(b"stub")
    (nested / "b.xlsx").write_bytes(b"stub")
    (nested / "c.csv").write_text("name;qty\nКабель;120\n", encoding="utf-8")
    (nested / "ignore.txt").write_text("x", encoding="utf-8")

    files = _discover_documents(root)

    assert [path.name for path in files] == ["a.docx", "b.xlsx", "c.csv"]


def test_extract_csv_text_includes_rows(tmp_path: Path) -> None:
    path = tmp_path / "spec.csv"
    path.write_text("Наименование;Кол-во\nСветильник;12\n", encoding="utf-8")

    text = _extract_csv_text(path)

    assert "Наименование | Кол-во" in text
    assert "Светильник | 12" in text


def test_extract_docx_text_includes_paragraphs_and_table(tmp_path: Path) -> None:
    path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("Первая строка")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Кабель"
    table.rows[0].cells[1].text = "120 м"
    doc.save(path)

    text = _extract_docx_text(path)

    assert "Первая строка" in text
    assert "Кабель | 120 м" in text


def test_extract_xlsx_text_includes_sheet_name_and_cells(tmp_path: Path) -> None:
    path = tmp_path / "spec.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "СО"
    sheet.append(["Наименование", "Кол-во"])
    sheet.append(["Светильник", 12])
    workbook.save(path)

    text = _extract_xlsx_text(path)

    assert "# СО" in text
    assert "Наименование | Кол-во" in text
    assert "Светильник | 12" in text


def test_extract_glossary_terms_from_pages_splits_articles() -> None:
    pages = [
        (
            5,
            """
            ГОСТ Р 58033—2017
            3.1.1 объект (капитального) строительства (construction works): Здание,
            строение, сооружение.
            3.1.2 сооружение (civil engineering works): Объекты завершенного строительства.
            """,
        )
    ]

    terms = _extract_glossary_terms_from_pages(pages)

    assert [term.article_no for term in terms] == ["3.1.1", "3.1.2"]
    assert terms[0].term == "объект (капитального) строительства"
    assert terms[0].english_term == "construction works"
    assert terms[0].definition == "Здание, строение, сооружение."
    assert terms[0].page == 5


def test_extract_glossary_terms_from_pages_ignores_headings_and_appends_multiline_text() -> None:
    pages = [
        (
            5,
            """
            3 Типы зданий и гражданских сооружений
            3.1 Основные термины
            3.1.3 здание (building): Объект, предназначенный для постоянного
            или временного пребывания в нем людей.
            3.2 Сооружения
            """,
        ),
        (
            6,
            """
            3.2.1 работы земляные (earthworks): Комплекс строительных работ,
            включающий выемку грунта.
            """,
        ),
    ]

    terms = _extract_glossary_terms_from_pages(pages)

    assert [term.article_no for term in terms] == ["3.1.3", "3.2.1"]
    assert terms[0].definition == (
        "Объект, предназначенный для постоянного или временного пребывания в нем людей."
    )
    assert terms[1].definition == "Комплекс строительных работ, включающий выемку грунта."


def test_compute_md5_hex_returns_32_char_hash(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    source.write_bytes(b"abc")

    digest = _compute_md5_hex(source)

    assert digest == "900150983cd24fb0d6963f7d28e17f72"

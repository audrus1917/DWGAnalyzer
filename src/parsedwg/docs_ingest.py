from __future__ import annotations

import asyncio
import csv
import hashlib
import re

from dataclasses import dataclass
from pathlib import Path

from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from .db import async_session_factory
from .orm import Entity, EntityType

SUPPORTED_DOC_SUFFIXES = {".pdf", ".docx", ".xlsx", ".csv"}

_GLOSSARY_ARTICLE_RE = re.compile(
    r"^(?P<article_no>\d+\.\d+\.\d+)\s+"
    r"(?P<term>.+)\s+\((?P<english_term>[^()]+)\):\s*(?P<definition>.*)$"
)
_SECTION_HEADING_RE = re.compile(r"^\d+(?:\.\d+)?\s+.+$")


@dataclass(frozen=True, slots=True)
class GlossaryTerm:
    article_no: str
    term: str
    english_term: str
    definition: str
    page: int


@dataclass(slots=True)
class _GlossaryTermBuilder:
    article_no: str
    term: str
    english_term: str
    definition_parts: list[str]
    page: int


def _compute_md5_hex(path: Path) -> str:
    digest = hashlib.md5(usedforsecurity=False)
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _discover_documents(source_path: Path) -> list[Path]:
    if not source_path.exists():
        raise FileNotFoundError(f"Путь {source_path} не найден.")

    if source_path.is_file():
        if source_path.suffix.lower() not in SUPPORTED_DOC_SUFFIXES:
            raise ValueError("Поддерживаются только PDF, DOCX, XLSX и CSV файлы.")
        return [source_path]

    files = sorted(
        item
        for item in source_path.rglob("*")
        if item.is_file() and item.suffix.lower() in SUPPORTED_DOC_SUFFIXES
    )
    if not files:
        raise ValueError(f"В каталоге {source_path} не найдено PDF/DOCX/XLSX/CSV файлов.")
    return files


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    chunks: list[str] = []

    chunks.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    )

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if values:
                chunks.append(" | ".join(values))

    return "\n".join(chunks)


def _extract_xlsx_text(path: Path) -> str:
    workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
    chunks: list[str] = []
    try:
        for sheet in workbook.worksheets:
            chunks.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    chunks.append(" | ".join(values))
    finally:
        workbook.close()
    return "\n".join(chunks)


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        normalized = text.strip()
        if normalized:
            chunks.append(normalized)
    return "\n\n".join(chunks)


def _extract_csv_text(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    sample = raw[:4096]
    delimiter = ","
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        if ";" in sample and "," not in sample:
            delimiter = ";"

    chunks: list[str] = []
    reader = csv.reader(raw.splitlines(), delimiter=delimiter)
    for row in reader:
        values = [cell.strip() for cell in row if cell and cell.strip()]
        if values:
            chunks.append(" | ".join(values))
    return "\n".join(chunks)


def _normalize_pdf_page_text(text: str) -> list[str]:
    normalized = text.replace("\x00", "").replace("\u00ad", "")
    normalized = re.sub(r"(?<=[A-Za-zА-Яа-я])-\n(?=[A-Za-zА-Яа-я])", "", normalized)
    normalized = re.sub(r"(?<=[A-Za-zА-Яа-я])\n(?=[a-zа-я])", "", normalized)
    return [line.strip() for line in normalized.splitlines() if line.strip()]


def _should_skip_pdf_line(line: str) -> bool:
    if re.fullmatch(r"\d+", line):
        return True
    if line.startswith("ГОСТ"):
        return True
    if line == "Издание официальное":
        return True
    return False


def _clean_definition_text(parts: list[str]) -> str:
    text = " ".join(part.strip() for part in parts if part.strip())
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _extract_glossary_terms_from_pages(pages: list[tuple[int, str]]) -> list[GlossaryTerm]:
    terms: list[GlossaryTerm] = []
    current: _GlossaryTermBuilder | None = None

    def flush_current() -> None:
        nonlocal current
        if current is None:
            return

        definition = _clean_definition_text(current.definition_parts)
        if definition:
            terms.append(
                GlossaryTerm(
                    article_no=current.article_no,
                    term=current.term,
                    english_term=current.english_term,
                    definition=definition,
                    page=current.page,
                )
            )
        current = None

    for page_num, page_text in pages:
        for line in _normalize_pdf_page_text(page_text):
            if _should_skip_pdf_line(line):
                continue

            match = _GLOSSARY_ARTICLE_RE.match(line)
            if match:
                flush_current()
                current = _GlossaryTermBuilder(
                    article_no=match.group("article_no"),
                    term=match.group("term").strip(),
                    english_term=match.group("english_term").strip(),
                    definition_parts=[match.group("definition").strip()],
                    page=page_num,
                )
                continue

            if current is None:
                continue

            if _SECTION_HEADING_RE.match(line):
                continue

            current.definition_parts.append(line)

    flush_current()
    return terms


def _extract_glossary_terms(path: Path) -> list[GlossaryTerm]:
    reader = PdfReader(str(path))
    pages = [
        (page_num, page.extract_text() or "")
        for page_num, page in enumerate(reader.pages, start=1)
    ]
    return _extract_glossary_terms_from_pages(pages)


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".xlsx":
        return _extract_xlsx_text(path)
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".csv":
        return _extract_csv_text(path)
    raise ValueError(f"Неподдерживаемый тип документа: {path.suffix}")


async def _save_documents_to_db(source_path: Path, documents: list[Path]) -> int:
    created = 0

    async with async_session_factory() as session:
        session: AsyncSession

        root = Entity(
            name=source_path.name if source_path.name else str(source_path),
            description="Корневая папка импортированных документов PDF/DOCX/XLSX/CSV",
            entity_type=EntityType.folder,
            data={"path": str(source_path)},
            start_from=str(source_path),
        )
        session.add(root)
        await session.flush()

        for doc_path in documents:
            text = _extract_text(doc_path)
            glossary_terms = _extract_glossary_terms(doc_path) if doc_path.suffix.lower() == ".pdf" else []
            try:
                rel_path = str(doc_path.relative_to(source_path))
            except ValueError:
                rel_path = doc_path.name

            entity = Entity(
                name=doc_path.name,
                description=text,
                entity_type=EntityType.file,
                data={
                    "doc_type": doc_path.suffix.lower().lstrip("."),
                    "relative_path": rel_path,
                    "size_bytes": doc_path.stat().st_size,
                },
                file_md5=_compute_md5_hex(doc_path),
                start_from=str(doc_path),
                parent_id=root.id,
            )
            session.add(entity)
            await session.flush()
            created += 1

            for term in glossary_terms:
                session.add(
                    Entity(
                        name=term.term,
                        description=term.definition,
                        entity_type=EntityType.primitive,
                        data={
                            "article_no": term.article_no,
                            "english_term": term.english_term,
                            "page": term.page,
                            "source_kind": "glossary_term",
                            "source_file": rel_path,
                        },
                        start_from=f"{doc_path}#page={term.page}",
                        parent_id=entity.id,
                    )
                )
                created += 1

        await session.commit()

    return created


def run_documents_ingest(source_path: Path) -> dict[str, object]:
    """Рекурсивно импортирует PDF/DOCX/XLSX/CSV документы в таблицу entity."""

    source = source_path.resolve()
    documents = _discover_documents(source)
    created = asyncio.run(_save_documents_to_db(source, documents))
    return {
        "doc_count": len(documents),
        "created_entities": created,
        "source": str(source),
    }


__all__ = [
    "GlossaryTerm",
    "SUPPORTED_DOC_SUFFIXES",
    "run_documents_ingest",
    "_extract_glossary_terms_from_pages",
]

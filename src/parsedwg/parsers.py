from __future__ import annotations

import re
import shutil
import subprocess

from deprecated import deprecated
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from ezdxf.addons.odafc import readfile as read_odafc

from ezdxf.filemanagement import readfile

from .models import ParsedItem

ITEM_RE = re.compile(
    r"(?P<name>.+?)\s*(?:-|—|–|;|:)\s*(?P<qty>\d+(?:[.,]\d+)?)\s*"
    r"(?P<unit>м2|м3|м|шт|компл\.?|кг|т|л)\b",
    re.IGNORECASE,
)


def classify_section(name: str) -> str:
    lowered = name.lower()
    if any(word in lowered for word in ("монтаж", "проклад", "установк", "демонтаж")):
        return "works"
    if any(word in lowered for word in ("светильник", "щит", "насос", "вентилятор", "датчик")):
        return "equipment"
    return "materials"


def parse_item_line(line: str, source: str = "text") -> ParsedItem | None:
    cleaned = " ".join(line.split())
    if not cleaned:
        return None

    match = ITEM_RE.search(cleaned)
    if not match:
        return None

    name = match.group("name").strip(" -–—;:,.\t")
    quantity = float(match.group("qty").replace(",", "."))
    unit = match.group("unit").lower().rstrip(".")
    return ParsedItem(
        name=name,
        quantity=quantity,
        unit=unit,
        section=classify_section(name),
        source=source,
    )


def parse_text_blob(text: str, source: str) -> list[ParsedItem]:
    items: list[ParsedItem] = []
    for line in text.splitlines():
        item = parse_item_line(line, source=source)
        if item is not None:
            items.append(item)
    return items


def _extract_dxf_text(path: Path) -> str:
    doc = readfile(str(path))
    lines: list[str] = []

    for entity in doc.modelspace():
        entity_type = entity.dxftype()
        if entity_type == "TEXT":
            value = entity.dxf.text.strip()
            if value:
                lines.append(value)
        elif entity_type == "MTEXT":
            plain_text = getattr(entity, "plain_text", None)
            if callable(plain_text):
                value = str(plain_text()).strip()
                if value:
                    lines.append(value)
        elif entity_type == "INSERT":
            for attrib in getattr(entity, "attribs", []):
                value = attrib.dxf.text.strip()
                if value:
                    lines.append(value)

    return "\n".join(lines)


def convert_dwg(path: Path) -> Path:
    """Convert a DWG file to DXF and return the temporary DXF path.

    Args:
        path: Path to the DWG file.

    Returns:
        Path to the created temporary DXF file.

    Raises:
        RuntimeError: If ODA File Converter is not available in PATH.
    """

    converter = shutil.which("ODAFileConverter") or shutil.which("odafc")
    if converter is None:
        raise RuntimeError(
            "DWG processing requires ODA File Converter in PATH. "
            "Alternatively, load the DXF file directly."
        )

    converted = read_odafc(path)
    target = path.with_suffix(".converted.dxf")
    converted.saveas(target)
    return target


def parse_drawing_file(path: str | Path) -> list[ParsedItem]:
    """Parse a DWG/DXF file and extract items.

    Args:
        path: Path to the drawing file.

    Returns:
        List of extracted items.

    Raises:
        RuntimeError: If ODA File Converter is unavailable for DWG.
        ValueError: If the file format is unsupported.
    """
    source_path = Path(path)
    suffix = source_path.suffix.lower()
    if suffix == ".dwg":
        dxf_path = convert_dwg(source_path)
        text = _extract_dxf_text(dxf_path)
    elif suffix in {".dxf", ".dxb"}:
        text = _extract_dxf_text(source_path)
    else:
        raise ValueError("Only DWG and DXF files are supported.")

    return parse_text_blob(text, source=source_path.suffix.lower().lstrip("."))


def parse_note_file(path: str | Path) -> list[ParsedItem]:
    """Parse a text or DOCX explanatory note.

    Args:
        path: Path to the note file.

    Returns:
        List of extracted items.

    Raises:
        ValueError: If the file format is unsupported.
    """
    source_path = Path(path)
    suffix = source_path.suffix.lower()

    if suffix in {".txt", ".md", ".rst"}:
        text = source_path.read_text(encoding="utf-8")
    elif suffix == ".docx":
        doc = Document(str(source_path))
        lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    lines.append(" - ".join(values))
        text = "\n".join(lines)
    else:
        raise ValueError("Only TXT, RST, MD, and DOCX explanatory notes are supported.")

    return parse_text_blob(text, source=source_path.suffix.lower().lstrip("."))
